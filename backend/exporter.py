"""Export della sessione in TXT / Markdown / DOCX / PDF / JSON."""
import io
import json
import re
from datetime import datetime
from typing import Tuple

from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches


LEVEL_LABELS = {
    "bambino": "Bambino (10 anni)",
    "superiori": "Scuole superiori",
    "universitario": "Universitario",
    "esperto": "Esperto",
    "tecnico": "Tecnico",
    "semplice": "Semplice",
}


def _safe_filename(name: str) -> str:
    """Sanitize a string to be used as a filename."""
    s = re.sub(r"[^\w\-. ]+", "_", (name or "lezione").strip())
    return s[:80] or "lezione"


def _strip_md(text: str) -> str:
    """Rimuove markdown base per output plain text."""
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    return text


def _md_inline_to_docx_run(paragraph, text: str, base_bold: bool = False):
    """Convert inline **bold** and *italic* to python-docx runs."""
    # Simple state machine: split by ** first, then within by *
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        else:
            # Handle italic *x*
            sub_parts = re.split(r"(\*[^*]+\*)", part)
            for sp in sub_parts:
                if sp.startswith("*") and sp.endswith("*") and len(sp) > 2:
                    r = paragraph.add_run(sp[1:-1])
                    r.italic = True
                elif sp:
                    r = paragraph.add_run(sp)
                    if base_bold:
                        r.bold = True


def build_txt(session: dict) -> bytes:
    a = session.get("analysis", {})
    lines = []
    lines.append("=" * 70)
    lines.append(f"INTELIGENT STUDY — LEZIONE")
    lines.append("=" * 70)
    lines.append("")
    lines.append(f"Titolo: {a.get('topic', session.get('title', ''))}")
    lines.append(f"Livello: {LEVEL_LABELS.get(session.get('level'), session.get('level', ''))}")
    lines.append(f"Difficoltà: {a.get('difficulty', '-')}")
    lines.append(f"Durata stimata: {a.get('estimated_minutes', '-')} minuti")
    lines.append(f"Fonte: {session.get('source', {}).get('title', '-')}")
    lines.append(f"Creata: {session.get('created_at', '')}")
    lines.append("")

    if a.get("summary"):
        lines.append("RIASSUNTO")
        lines.append("-" * 70)
        lines.append(a["summary"])
        lines.append("")

    if a.get("prerequisites"):
        lines.append("PREREQUISITI")
        lines.append("-" * 70)
        for p in a["prerequisites"]:
            lines.append(f"  • {p}")
        lines.append("")

    if a.get("concepts"):
        lines.append("CONCETTI CHIAVE")
        lines.append("-" * 70)
        for c in a["concepts"]:
            lines.append(f"  • {c.get('name', '')}: {c.get('definition', '')}")
        lines.append("")

    if a.get("chapters"):
        lines.append("CAPITOLI")
        lines.append("-" * 70)
        for i, c in enumerate(a["chapters"], 1):
            lines.append(f"  {i}. {c.get('title', '')}")
            if c.get("objective"):
                lines.append(f"     Obiettivo: {c['objective']}")
            for kp in c.get("key_points", []):
                lines.append(f"     - {kp}")
            lines.append("")

    if a.get("formulas"):
        lines.append("FORMULE")
        lines.append("-" * 70)
        for f in a["formulas"]:
            lines.append(f"  • {f}")
        lines.append("")

    if a.get("examples"):
        lines.append("ESEMPI")
        lines.append("-" * 70)
        for e in a["examples"]:
            lines.append(f"  • {e}")
        lines.append("")

    msgs = session.get("messages", [])
    if msgs:
        lines.append("CONVERSAZIONE CON IL MAESTRO")
        lines.append("=" * 70)
        for m in msgs:
            role = "TU" if m.get("role") == "user" else "MAESTRO"
            lines.append("")
            lines.append(f"[{role}]")
            lines.append(_strip_md(m.get("content", "")))
        lines.append("")

    quiz = session.get("quiz")
    if quiz:
        lines.append("=" * 70)
        lines.append("QUIZ")
        lines.append("=" * 70)
        for i, q in enumerate(quiz.get("multiple_choice", []), 1):
            lines.append(f"\n{i}. {q.get('question', '')}")
            for j, opt in enumerate(q.get("options", [])):
                marker = " ✓" if j == q.get("correct_index") else ""
                lines.append(f"   {chr(97 + j)}) {opt}{marker}")
            if q.get("explanation"):
                lines.append(f"   → {q['explanation']}")
        for q in quiz.get("true_false", []):
            lines.append(f"\n{q.get('statement', '')}")
            lines.append(f"   Risposta: {'Vero' if q.get('answer') else 'Falso'}")
            if q.get("explanation"):
                lines.append(f"   → {q['explanation']}")
        for q in quiz.get("open_questions", []):
            lines.append(f"\nDomanda aperta: {q.get('question', '')}")
            lines.append(f"   Risposta modello: {q.get('sample_answer', '')}")
        if quiz.get("flashcards"):
            lines.append("\n--- FLASHCARD ---")
            for c in quiz["flashcards"]:
                lines.append(f"  Q: {c.get('front', '')}")
                lines.append(f"  A: {c.get('back', '')}")

    notes = session.get("notes", [])
    if notes:
        lines.append("")
        lines.append("=" * 70)
        lines.append("LE MIE NOTE")
        lines.append("=" * 70)
        for n in notes:
            lines.append(f"  • {n.get('text', '')}")

    lines.append("")
    lines.append("-" * 70)
    lines.append("Generato da inteligent STUDY · Maestro AI")
    return "\n".join(lines).encode("utf-8")


def build_markdown(session: dict) -> bytes:
    a = session.get("analysis", {})
    out = []
    out.append(f"# {a.get('topic', session.get('title', 'Lezione'))}")
    out.append("")
    out.append(f"> **Livello**: {LEVEL_LABELS.get(session.get('level'), session.get('level', ''))} · "
               f"**Difficoltà**: {a.get('difficulty', '-')} · "
               f"**Durata**: {a.get('estimated_minutes', '-')} min")
    out.append(f"> **Fonte**: {session.get('source', {}).get('title', '-')}")
    out.append("")

    if a.get("summary"):
        out.append("## Riassunto")
        out.append(a["summary"])
        out.append("")

    if a.get("prerequisites"):
        out.append("## Prerequisiti")
        for p in a["prerequisites"]:
            out.append(f"- {p}")
        out.append("")

    if a.get("concepts"):
        out.append("## Concetti chiave")
        for c in a["concepts"]:
            out.append(f"- **{c.get('name', '')}** — {c.get('definition', '')}")
        out.append("")

    if a.get("chapters"):
        out.append("## Capitoli")
        for i, c in enumerate(a["chapters"], 1):
            out.append(f"### {i}. {c.get('title', '')}")
            if c.get("objective"):
                out.append(f"*Obiettivo*: {c['objective']}")
            for kp in c.get("key_points", []):
                out.append(f"- {kp}")
            out.append("")

    msgs = session.get("messages", [])
    if msgs:
        out.append("## Lezione (conversazione)")
        out.append("")
        for m in msgs:
            role = "**Tu**" if m.get("role") == "user" else "**Maestro**"
            out.append(f"---")
            out.append(f"{role}:")
            out.append("")
            out.append(m.get("content", ""))
            out.append("")

    quiz = session.get("quiz")
    if quiz:
        out.append("## Quiz")
        for i, q in enumerate(quiz.get("multiple_choice", []), 1):
            out.append(f"**{i}. {q.get('question', '')}**")
            for j, opt in enumerate(q.get("options", [])):
                marker = " ✅" if j == q.get("correct_index") else ""
                out.append(f"- {chr(97 + j)}) {opt}{marker}")
            if q.get("explanation"):
                out.append(f"> {q['explanation']}")
            out.append("")

    notes = session.get("notes", [])
    if notes:
        out.append("## Le mie note")
        for n in notes:
            out.append(f"- {n.get('text', '')}")

    out.append("")
    out.append("---")
    out.append("*Generato da inteligent STUDY · Maestro AI*")
    return "\n".join(out).encode("utf-8")


def build_json(session: dict) -> bytes:
    return json.dumps(session, ensure_ascii=False, indent=2).encode("utf-8")


def build_docx(session: dict) -> bytes:
    a = session.get("analysis", {})
    doc = Document()

    # Styles
    styles = doc.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(11)

    # Title
    h = doc.add_heading(a.get("topic", session.get("title", "Lezione")), level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x06, 0x5F, 0x46)

    # Meta
    meta = doc.add_paragraph()
    meta.add_run("Livello: ").bold = True
    meta.add_run(f"{LEVEL_LABELS.get(session.get('level'), session.get('level', ''))}    ")
    meta.add_run("Difficoltà: ").bold = True
    meta.add_run(f"{a.get('difficulty', '-')}    ")
    meta.add_run("Durata stimata: ").bold = True
    meta.add_run(f"{a.get('estimated_minutes', '-')} min")

    src = session.get("source", {})
    if src.get("title"):
        p = doc.add_paragraph()
        p.add_run("Fonte: ").bold = True
        p.add_run(src.get("title", ""))

    if a.get("summary"):
        doc.add_heading("Riassunto", level=1)
        doc.add_paragraph(a["summary"])

    if a.get("prerequisites"):
        doc.add_heading("Prerequisiti", level=1)
        for p in a["prerequisites"]:
            doc.add_paragraph(p, style="List Bullet")

    if a.get("concepts"):
        doc.add_heading("Concetti chiave", level=1)
        for c in a["concepts"]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"{c.get('name', '')}: ")
            r.bold = True
            p.add_run(c.get("definition", ""))

    if a.get("chapters"):
        doc.add_heading("Capitoli", level=1)
        for i, c in enumerate(a["chapters"], 1):
            doc.add_heading(f"{i}. {c.get('title', '')}", level=2)
            if c.get("objective"):
                p = doc.add_paragraph()
                r = p.add_run("Obiettivo: ")
                r.bold = True
                r.italic = True
                p.add_run(c["objective"])
            for kp in c.get("key_points", []):
                doc.add_paragraph(kp, style="List Bullet")

    if a.get("formulas"):
        doc.add_heading("Formule", level=1)
        for f in a["formulas"]:
            doc.add_paragraph(f, style="List Bullet")

    if a.get("examples"):
        doc.add_heading("Esempi", level=1)
        for e in a["examples"]:
            doc.add_paragraph(e, style="List Bullet")

    msgs = session.get("messages", [])
    if msgs:
        doc.add_page_break()
        doc.add_heading("Conversazione con il Maestro", level=1)
        for m in msgs:
            role = "Tu" if m.get("role") == "user" else "Maestro"
            head = doc.add_paragraph()
            r = head.add_run(role)
            r.bold = True
            r.font.color.rgb = RGBColor(0x06, 0x5F, 0x46) if role == "Maestro" else RGBColor(0x33, 0x33, 0x33)
            content = m.get("content", "")
            for line in content.split("\n"):
                if line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("- ") or line.startswith("* "):
                    p = doc.add_paragraph(style="List Bullet")
                    _md_inline_to_docx_run(p, line[2:])
                elif re.match(r"^\d+\.\s", line):
                    p = doc.add_paragraph(style="List Number")
                    _md_inline_to_docx_run(p, re.sub(r"^\d+\.\s", "", line))
                elif line.strip():
                    p = doc.add_paragraph()
                    _md_inline_to_docx_run(p, line)
                else:
                    doc.add_paragraph("")

    quiz = session.get("quiz")
    if quiz:
        doc.add_page_break()
        doc.add_heading("Quiz", level=1)
        for i, q in enumerate(quiz.get("multiple_choice", []), 1):
            p = doc.add_paragraph()
            r = p.add_run(f"{i}. {q.get('question', '')}")
            r.bold = True
            for j, opt in enumerate(q.get("options", [])):
                marker = "  ✓" if j == q.get("correct_index") else ""
                doc.add_paragraph(f"{chr(97 + j)}) {opt}{marker}", style="List Bullet")
            if q.get("explanation"):
                p = doc.add_paragraph()
                r = p.add_run(q["explanation"])
                r.italic = True

        for q in quiz.get("true_false", []):
            p = doc.add_paragraph()
            r = p.add_run(q.get("statement", ""))
            r.bold = True
            p.add_run(f"  → {'Vero' if q.get('answer') else 'Falso'}")
            if q.get("explanation"):
                p2 = doc.add_paragraph()
                r2 = p2.add_run(q["explanation"])
                r2.italic = True

        for q in quiz.get("open_questions", []):
            p = doc.add_paragraph()
            r = p.add_run(q.get("question", ""))
            r.bold = True
            doc.add_paragraph(f"Risposta modello: {q.get('sample_answer', '')}")

        if quiz.get("flashcards"):
            doc.add_heading("Flashcard", level=2)
            for c in quiz["flashcards"]:
                p = doc.add_paragraph()
                r = p.add_run(f"Q: ")
                r.bold = True
                p.add_run(c.get("front", ""))
                p2 = doc.add_paragraph()
                r2 = p2.add_run(f"A: ")
                r2.bold = True
                p2.add_run(c.get("back", ""))

    notes = session.get("notes", [])
    if notes:
        doc.add_page_break()
        doc.add_heading("Le mie note", level=1)
        for n in notes:
            doc.add_paragraph(n.get("text", ""), style="List Bullet")

    doc.add_paragraph("")
    footer = doc.add_paragraph()
    r = footer.add_run("Generato da inteligent STUDY · Maestro AI")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class _StudyPDF(FPDF):
    def __init__(self):
        super().__init__(orientation="P", unit="mm", format="A4")
        self.set_auto_page_break(auto=True, margin=15)
        self.set_margins(15, 15, 15)
        # Use built-in Helvetica (supports basic latin, no Italian accents issue as we normalize)

    def header(self):
        if self.page_no() == 1:
            return
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(128)
        self.cell(0, 6, "inteligent STUDY · Maestro AI", ln=0, align="L")
        self.cell(0, 6, f"Pagina {self.page_no()}", ln=1, align="R")
        self.ln(2)

    def footer(self):
        pass


def _sanitize_pdf(text: str) -> str:
    """Replace unicode chars unsupported by latin-1 font."""
    if not text:
        return ""
    replacements = {
        "—": "-", "–": "-", "…": "...",
        "'": "'", "'": "'", """: '"', """: '"',
        "•": "-", "→": "->", "←": "<-", "↔": "<->",
        "✓": "[ok]", "✅": "[ok]", "✗": "[x]", "❌": "[x]",
        "✦": "*", "★": "*", "☆": "*",
        "🌱": "", "🔥": "", "🎓": "", "🧠": "", "📚": "", "🎯": "",
        "\u00a0": " ", "\u200b": "",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode("latin-1", "replace").decode("latin-1")


def build_pdf(session: dict) -> bytes:
    a = session.get("analysis", {})
    pdf = _StudyPDF()
    pdf.add_page()

    def mc(text, height=6, style="", size=11, color=(0, 0, 0)):
        """Safe multi_cell with position reset + CHAR wrap for long tokens."""
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", style, size)
        pdf.set_text_color(*color)
        try:
            pdf.multi_cell(0, height, _sanitize_pdf(text), wrapmode="CHAR", new_x="LMARGIN", new_y="NEXT")
        except Exception:
            # Ultima risorsa: split e re-render riga per riga
            for line in _sanitize_pdf(text).split("\n"):
                if line.strip():
                    pdf.set_x(pdf.l_margin)
                    pdf.multi_cell(0, height, line[:200], wrapmode="CHAR", new_x="LMARGIN", new_y="NEXT")

    # Title
    mc(a.get("topic", session.get("title", "Lezione")), height=10, style="B", size=22, color=(6, 95, 70))
    pdf.ln(1)

    mc(
        f"Livello: {LEVEL_LABELS.get(session.get('level'), session.get('level', ''))}  |  "
        f"Difficolta: {a.get('difficulty', '-')}  |  "
        f"Durata: {a.get('estimated_minutes', '-')} min",
        height=5, style="I", size=10, color=(100, 100, 100),
    )
    src = session.get("source", {})
    if src.get("title"):
        mc(f"Fonte: {src.get('title', '')}", height=5, style="I", size=10, color=(100, 100, 100))
    pdf.ln(4)

    def h(text, level=1):
        pdf.ln(2)
        mc(text, height=8, style="B", size=16 if level == 1 else 13, color=(6, 95, 70))

    def bullets(items):
        for it in items:
            mc(f"  -  {it}", height=6)

    if a.get("summary"):
        h("Riassunto"); mc(a["summary"])

    if a.get("prerequisites"):
        h("Prerequisiti"); bullets(a["prerequisites"])

    if a.get("concepts"):
        h("Concetti chiave")
        for c in a["concepts"]:
            mc(f"  - {c.get('name', '')}", style="B")
            mc(f"     {c.get('definition', '')}")
            pdf.ln(1)

    if a.get("chapters"):
        h("Capitoli")
        for i, c in enumerate(a["chapters"], 1):
            h(f"{i}. {c.get('title', '')}", level=2)
            if c.get("objective"):
                mc(f"Obiettivo: {c['objective']}", height=5, style="I", size=10)
            bullets(c.get("key_points", []))

    if a.get("formulas"):
        h("Formule"); bullets(a["formulas"])

    if a.get("examples"):
        h("Esempi"); bullets(a["examples"])

    msgs = session.get("messages", [])
    if msgs:
        pdf.add_page()
        h("Conversazione con il Maestro")
        for m in msgs:
            role = "Tu" if m.get("role") == "user" else "Maestro"
            mc(role, style="B", color=(6, 95, 70) if role == "Maestro" else (40, 40, 40))
            mc(_strip_md(m.get("content", "")))
            pdf.ln(2)

    quiz = session.get("quiz")
    if quiz:
        pdf.add_page()
        h("Quiz")
        for i, q in enumerate(quiz.get("multiple_choice", []), 1):
            mc(f"{i}. {q.get('question', '')}", style="B")
            for j, opt in enumerate(q.get("options", [])):
                marker = "  [corretta]" if j == q.get("correct_index") else ""
                mc(f"  {chr(97 + j)}) {opt}{marker}")
            if q.get("explanation"):
                mc(f"  -> {q['explanation']}", height=5, style="I", size=10, color=(80, 80, 80))
            pdf.ln(2)
        for q in quiz.get("true_false", []):
            mc(q.get("statement", ""), style="B")
            mc(f"  Risposta: {'Vero' if q.get('answer') else 'Falso'}")
            pdf.ln(1)
        for q in quiz.get("open_questions", []):
            mc(q.get("question", ""), style="B")
            mc(f"  Risposta modello: {q.get('sample_answer', '')}")
            pdf.ln(1)
        if quiz.get("flashcards"):
            h("Flashcard", level=2)
            for c in quiz["flashcards"]:
                mc(f"Q: {c.get('front', '')}", style="B")
                mc(f"A: {c.get('back', '')}")
                pdf.ln(1)

    notes = session.get("notes", [])
    if notes:
        pdf.add_page()
        h("Le mie note")
        bullets([n.get("text", "") for n in notes])

    pdf.ln(6)
    mc("Generato da inteligent STUDY - Maestro AI", height=4, style="I", size=8, color=(128, 128, 128))

    out = pdf.output(dest="S")
    return bytes(out) if isinstance(out, (bytes, bytearray)) else out.encode("latin-1")


FORMAT_CONFIG = {
    "txt": ("text/plain; charset=utf-8", "txt", build_txt),
    "md": ("text/markdown; charset=utf-8", "md", build_markdown),
    "json": ("application/json; charset=utf-8", "json", build_json),
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx", build_docx),
    "pdf": ("application/pdf", "pdf", build_pdf),
}


def export_session(session: dict, fmt: str) -> Tuple[bytes, str, str]:
    """Returns (data, content_type, filename)."""
    fmt = fmt.lower()
    if fmt not in FORMAT_CONFIG:
        raise ValueError(f"Formato non supportato: {fmt}")
    ct, ext, builder = FORMAT_CONFIG[fmt]
    data = builder(session)
    topic = session.get("analysis", {}).get("topic") or session.get("title") or "lezione"
    fname = f"{_safe_filename(topic)}.{ext}"
    return data, ct, fname
